import io
import fitz
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def convert_pdf_to_word(file_path):
    """
    Convert PDF to Word document by extracting text with formatting
    and images from each page.
    """
    pdf_doc = fitz.open(file_path)
    word_doc = Document()

    for page_num in range(pdf_doc.page_count):
        page = pdf_doc[page_num]

        if page_num > 0:
            word_doc.add_page_break()

        # Extract text with formatting info
        text_dict = page.get_text("dict")

        for block in text_dict.get("blocks", []):
            if block["type"] == 0:  # Text block
                for line in block.get("lines", []):
                    para = word_doc.add_paragraph()

                    for span in line.get("spans", []):
                        run = para.add_run(span.get("text", ""))

                        # Apply font size
                        font_size = span.get("size", 12)
                        run.font.size = Pt(font_size)

                        # Apply font name
                        font_name = span.get("font", "")
                        if font_name:
                            # Map PDF font names to common fonts
                            if "bold" in font_name.lower():
                                run.bold = True
                            if "italic" in font_name.lower() or "oblique" in font_name.lower():
                                run.italic = True

                            # Clean font name
                            clean_name = font_name.split("-")[0].split("+")[-1]
                            clean_name = clean_name.replace("Bold", "").replace("Italic", "").strip()
                            if clean_name:
                                run.font.name = clean_name

                        # Apply color
                        color_int = span.get("color", 0)
                        if color_int and color_int != 0:
                            r = (color_int >> 16) & 0xFF
                            g = (color_int >> 8) & 0xFF
                            b = color_int & 0xFF
                            run.font.color.rgb = RGBColor(r, g, b)

                        # Apply flags
                        flags = span.get("flags", 0)
                        if flags & 2 ** 0:  # superscript
                            run.font.superscript = True
                        if flags & 2 ** 1:  # italic
                            run.italic = True
                        if flags & 2 ** 4:  # bold
                            run.bold = True

            elif block["type"] == 1:  # Image block
                try:
                    img_data = block.get("image")
                    if img_data:
                        img_stream = io.BytesIO(img_data)
                        # Calculate width based on block dimensions
                        width = block.get("width", 200)
                        # Convert PDF points to inches (72 points = 1 inch)
                        width_inches = min(width / 72, 6)  # Max 6 inches
                        word_doc.add_picture(img_stream, width=Inches(width_inches))
                except Exception:
                    pass  # Skip images that fail to extract

    # Save to bytes
    output = io.BytesIO()
    word_doc.save(output)
    output.seek(0)
    pdf_doc.close()

    return output.getvalue()
