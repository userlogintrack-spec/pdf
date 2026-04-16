import io
from docx import Document


def convert_word_to_html(file_path):
    """
    Convert Word document to HTML with basic formatting.
    """
    doc = Document(file_path)
    html_parts = [
        '<!DOCTYPE html>',
        '<html lang="en">',
        '<head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0">',
        '<title>Converted Document</title>',
        '<style>',
        'body { font-family: Arial, sans-serif; max-width: 900px; margin: 0 auto; padding: 20px; line-height: 1.6; }',
        'table { border-collapse: collapse; width: 100%; margin: 10px 0; }',
        'td, th { border: 1px solid #ddd; padding: 8px; }',
        'th { background-color: #f2f2f2; }',
        '</style>',
        '</head>',
        '<body>',
    ]

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            html_parts.append('<br/>')
            continue

        # Escape HTML
        text_escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        style_name = para.style.name.lower() if para.style else ""

        if 'heading 1' in style_name:
            html_parts.append(f'<h1>{text_escaped}</h1>')
        elif 'heading 2' in style_name:
            html_parts.append(f'<h2>{text_escaped}</h2>')
        elif 'heading 3' in style_name:
            html_parts.append(f'<h3>{text_escaped}</h3>')
        elif 'heading' in style_name:
            html_parts.append(f'<h4>{text_escaped}</h4>')
        elif 'list' in style_name:
            html_parts.append(f'<li>{text_escaped}</li>')
        else:
            # Build with run-level formatting
            runs_html = []
            for run in para.runs:
                run_text = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                if not run_text:
                    continue
                if run.bold and run.italic:
                    runs_html.append(f'<strong><em>{run_text}</em></strong>')
                elif run.bold:
                    runs_html.append(f'<strong>{run_text}</strong>')
                elif run.italic:
                    runs_html.append(f'<em>{run_text}</em>')
                elif run.underline:
                    runs_html.append(f'<u>{run_text}</u>')
                else:
                    runs_html.append(run_text)

            if runs_html:
                html_parts.append(f'<p>{"".join(runs_html)}</p>')
            else:
                html_parts.append(f'<p>{text_escaped}</p>')

    # Handle tables
    for table in doc.tables:
        html_parts.append('<table>')
        for i, row in enumerate(table.rows):
            html_parts.append('<tr>')
            tag = 'th' if i == 0 else 'td'
            for cell in row.cells:
                cell_text = cell.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                html_parts.append(f'<{tag}>{cell_text}</{tag}>')
            html_parts.append('</tr>')
        html_parts.append('</table>')

    html_parts.append('</body></html>')
    return "\n".join(html_parts).encode("utf-8")
