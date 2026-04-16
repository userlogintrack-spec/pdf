from docx import Document


def convert_word_to_text(file_path):
    """
    Extract all text from a Word document.
    """
    doc = Document(file_path)
    text_parts = []

    for para in doc.paragraphs:
        text_parts.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        text_parts.append("")  # Blank line before table
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            text_parts.append(row_text)

    return "\n".join(text_parts).encode("utf-8")
